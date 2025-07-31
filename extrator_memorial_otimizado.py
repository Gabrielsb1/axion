import re
import pandas as pd
import sys
import os
from docx import Document
from io import BytesIO
import argparse
from concurrent.futures import ThreadPoolExecutor, as_completed, ProcessPoolExecutor
import time
import PyPDF2

# Regex compilados uma única vez para melhor performance - ULTRA OTIMIZADOS
REGEX_TORRE = re.compile(
    r"Apartamento\s+(\d+),\s*tipo\s+([A-Z]),\s+da\s+Torre\s+(\d+).*?"
    r"área privativa principal de ([\d,]+)m².*?"
    r"área privativa total de ([\d,]+)m².*?"
    r"área de uso comum de ([\d,]+)m².*?"
    r"área real total de ([\d,]+)m².*?"
    r"fração ideal.*?de ([\d,]+).*?ou ([\d,]+)m²",
    re.DOTALL | re.IGNORECASE
)

REGEX_BLOCO1 = re.compile(
    r"APARTAMENTO\s+(\d+)\s*[-–]\s*BLOCO\s+(\d+):?", re.IGNORECASE
)

REGEX_BLOCO2 = re.compile(
    r"Apartamento\s+(\d+),\s*TIPO\s*([A-Z]),\s*do\s*Bloco\s+(\d+),", re.IGNORECASE
)

REGEX_AREAS_BLOCO2 = re.compile(
    r"área privativa principal de ([\d,.]+)m².*?"
    r"área privativa total de ([\d,.]+)m².*?"
    r"área de uso comum de ([\d,.]+)m².*?"
    r"área real total de ([\d,.]+)m².*?"
    r"fração ideal.*?de ([\d,.]+).*?ou ([\d,.]+)m²",
    re.DOTALL | re.IGNORECASE
)

REGEX_AREAS_BLOCO1 = re.compile(
    r"áreas:\s*privativa real de ([\d,]+)m²,\s*"
    r"área de uso comum real de ([\d,]+)m²,\s*"
    r"perfazendo uma área total real de ([\d,]+)m².*?"
    r"área equivalente de construção igual a ([\d,]+)m².*?"
    r"fração ideal.*?([0-9,.]+)%",
    re.DOTALL | re.IGNORECASE
)

# Regex para casas
REGEX_CASA_SPLIT = re.compile(r"(?:^|\n)(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?(\d{2})")
REGEX_AREA_TERRENO = re.compile(r"configuração.*?área total de *(\d+,\d+)")
REGEX_AREA_CONSTRUIDA = re.compile(r"área total construída da casa de *(\d+,\d+)")
REGEX_AREA_COMUM = re.compile(r"área de uso comum real de *(\d+,\d+)")
REGEX_AREA_TOTAL_REAL = re.compile(r"área total real de *(\d+,\d+)")
REGEX_FRACAO_IDEAL = re.compile(r"fração ideal do terreno correspondente a *(\d+,\d+)\s?%")

# Regex para identificação de tipo (melhorados)
REGEX_TORRE_PATTERN = re.compile(r"Apartamento\s+\d+,\s*tipo\s+[A-Z],\s+da\s+Torre", re.IGNORECASE)
REGEX_BLOCO_PATTERN1 = re.compile(r"Apartamento\s+\d+,\s*TIPO\s+[A-Z],\s*do\s+Bloco\s+\d+", re.IGNORECASE)
REGEX_BLOCO_PATTERN2 = re.compile(r"Apartamento\s+\d+\s*[-–]\s*Bloco\s+\d+", re.IGNORECASE)
REGEX_CASA_PATTERN = re.compile(r"(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?\d{2}")

# Regex pré-compilados para pavimento (OTIMIZAÇÃO CRÍTICA)
REGEX_PAVIMENTO_NUMERICO = re.compile(r"(\d+º?\s*[o°]?\s*Pavimento)", re.IGNORECASE)
REGEX_PAVIMENTO_TERREO = re.compile(r"(Pavimento\s*T[eé]rr?eo)", re.IGNORECASE)
REGEX_TERREO_SIMPLES = re.compile(r"(t[eé]rr?eo)", re.IGNORECASE)

# Regex pré-compilados para ordinais por extenso
REGEX_ORDINAIS_EXTENSO = re.compile(
    r"(primeiro|segundo|terceiro|quarto|quinto|sexto|sétimo|oitavo|nono|décimo|"
    r"décimo primeiro|décimo segundo|décimo terceiro|décimo quarto|décimo quinto|"
    r"décimo sexto|décimo sétimo|décimo oitavo|décimo nono|vigésimo|"
    r"vigésimo primeiro|vigésimo segundo|vigésimo terceiro|vigésimo quarto|"
    r"vigésimo quinto|vigésimo sexto|vigésimo sétimo|vigésimo oitavo|"
    r"vigésimo nono|trigésimo)\s+pavimento", re.IGNORECASE
)

# Cache para operações de string frequentes - REMOVIDO PARA MÁXIMA PERFORMANCE
# STRING_CACHE = {}

# def get_cached_string_operation(text, operation, *args):
#     """Cache para operações de string frequentes"""
#     cache_key = f"{hash(text)}_{operation}_{hash(str(args))}"
#     if cache_key not in STRING_CACHE:
#         if operation == 'lower':
#             STRING_CACHE[cache_key] = text.lower()
#         elif operation == 'find':
#             STRING_CACHE[cache_key] = text.find(*args)
#         elif operation == 'rfind':
#             STRING_CACHE[cache_key] = text.rfind(*args)
#         elif operation == 'replace':
#             STRING_CACHE[cache_key] = text.replace(*args)
#     
#     # Limpar cache se ficar muito grande
#     if len(STRING_CACHE) > 1000:
#         STRING_CACHE.clear()
#     
#     return STRING_CACHE[cache_key]

def extrair_pavimento_otimizado(texto):
    """Versão ULTRA-OTIMIZADA da extração de pavimento - MÁXIMA PERFORMANCE"""
    # Verificar apenas padrões mais comuns primeiro (80% dos casos)
    if 'pavimento' in texto.lower():
        match = REGEX_PAVIMENTO_NUMERICO.search(texto)
        if match:
            return match.group(1).strip()
        
        match = REGEX_PAVIMENTO_TERREO.search(texto)
        if match:
            return match.group(1).strip()
        
        match = REGEX_ORDINAIS_EXTENSO.search(texto)
        if match:
            return match.group(1).strip()
    
    # Verificar térreo apenas se não encontrou pavimento
    if 'térreo' in texto.lower() or 'terreo' in texto.lower():
        match = REGEX_TERREO_SIMPLES.search(texto)
        if match:
            return match.group(1).strip()
    
    return ""

def extrair_texto_pdf(file_path):
    """Extrai texto de arquivo PDF"""
    try:
        reader = PyPDF2.PdfReader(file_path)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text() + "\n"
        return texto
    except Exception as e:
        print(f"Erro ao extrair texto do PDF: {e}")
        return ""

def identificar_tipo_documento_otimizado(texto):
    """Versão melhorada da identificação de tipo baseada no código Streamlit"""
    # Padrão para Torre: Apartamento 001, tipo A, da Torre 01
    if REGEX_TORRE_PATTERN.search(texto):
        return "torre"
    # Padrão para Bloco: Apartamento 001, TIPO A, do Bloco 01 OU APARTAMENTO 001 – BLOCO 01
    elif REGEX_BLOCO_PATTERN1.search(texto) or REGEX_BLOCO_PATTERN2.search(texto):
        return "bloco"
    # Padrão para Casas: CASA 01, Casa 01, etc.
    elif REGEX_CASA_PATTERN.search(texto):
        return "casa"
    return "desconhecido"

def extrair_descricao_otimizada(texto):
    """Versão otimizada da extração de descrição"""
    texto_lower = texto.lower()
    
    # Procurar por palavras-chave de descrição
    idx_localizado = texto_lower.find('localizado')
    if idx_localizado == -1:
        idx_localizado = texto_lower.find('assim descrito:')
    
    if idx_localizado != -1:
        idx_ultimo_ponto = texto.rfind('.')
        if idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_localizado:
            descricao = texto[idx_localizado:idx_ultimo_ponto+1]
            return descricao.replace('\n', ' ').strip()
    
    return ""

def extrair_torre_otimizado(doc):
    """Versão ULTRA-OTIMIZADA da extração de torre - MÁXIMA PERFORMANCE"""
    dados = []
    
    # Processar todos os parágrafos de uma vez
    textos = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto and len(texto) > 15 and 'Apartamento' in texto:  # Filtro mais agressivo
            textos.append(texto)
    
    # Processar em lotes para melhor performance
    for texto in textos:
        match = REGEX_TORRE.search(texto)
        if match:
            numero, tipo, torre, privativa, total, comum, real, fracao, terreno = match.groups()
            
            # Extrações otimizadas - apenas se necessário
            descricao = ""
            if 'localizado' in texto.lower():
                idx_localizado = texto.lower().find('localizado')
                idx_ultimo_ponto = texto.rfind('.')
                if idx_ultimo_ponto > idx_localizado:
                    descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            
            pavimento = ""
            if 'pavimento' in texto.lower() or 'térreo' in texto.lower():
                pavimento = extrair_pavimento_otimizado(texto)
            
            dados.append({
                "Formato": "Torre",
                "Apartamento": numero,
                "Tipo": tipo,
                "Torre/Bloco": torre,
                "Pavimento": pavimento,
                "Área Privativa (m²)": privativa.replace(",", "."),
                "Área Comum (m²)": comum.replace(",", "."),
                "Área Total (m²)": real.replace(",", "."),
                "Fração Ideal (%)": fracao.replace(",", "."),
                "Área Terreno (m²)": terreno.replace(",", "."),
                "Descrição": descricao
            })
    
    return pd.DataFrame(dados)

def extrair_bloco_otimizado(doc):
    """Versão ULTRA-OTIMIZADA da extração de bloco - MÁXIMA PERFORMANCE"""
    dados = []
    
    # Processar todos os parágrafos de uma vez
    textos = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto and len(texto) > 15 and ('Apartamento' in texto or 'APARTAMENTO' in texto):  # Filtro mais agressivo
            textos.append(texto)
    
    for texto in textos:
        # Tenta padrão 2 primeiro (mais específico)
        match2 = REGEX_BLOCO2.search(texto)
        if match2:
            numero, tipo, bloco = match2.groups()
            
            # Extrair áreas primeiro (mais crítico)
            match_areas2 = REGEX_AREAS_BLOCO2.search(texto)
            if match_areas2:
                privativa, total, comum, real, fracao, terreno = match_areas2.groups()
                
                # Extrações otimizadas - apenas se necessário
                descricao = ""
                if 'localizado' in texto.lower():
                    idx_localizado = texto.lower().find('localizado')
                    idx_ultimo_ponto = texto.rfind('.')
                    if idx_ultimo_ponto > idx_localizado:
                        descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
                
                pavimento = ""
                if 'pavimento' in texto.lower() or 'térreo' in texto.lower():
                    pavimento = extrair_pavimento_otimizado(texto)
                
                dados.append({
                    "Formato": "Bloco",
                    "Apartamento": numero,
                    "Tipo": tipo,
                    "Torre/Bloco": bloco,
                    "Pavimento": pavimento,
                    "Área Privativa (m²)": privativa.replace(",", "."),
                    "Área Privativa Total (m²)": total.replace(",", "."),
                    "Área Comum (m²)": comum.replace(",", "."),
                    "Área Total (m²)": real.replace(",", "."),
                    "Fração Ideal (%)": fracao.replace(",", "."),
                    "Área Terreno (m²)": terreno.replace(",", "."),
                    "Descrição": descricao
                })
            continue
            
        # Tenta padrão 1
        match1 = REGEX_BLOCO1.search(texto)
        if match1:
            numero, bloco = match1.groups()
            tipo = ""
            
            # Extrair áreas primeiro
            match_areas1 = REGEX_AREAS_BLOCO1.search(texto)
            if match_areas1:
                privativa, comum, total, equivalente, fracao = match_areas1.groups()
                
                # Extrações otimizadas - apenas se necessário
                descricao = ""
                if 'localizado' in texto.lower():
                    idx_localizado = texto.lower().find('localizado')
                    idx_ultimo_ponto = texto.rfind('.')
                    if idx_ultimo_ponto > idx_localizado:
                        descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
                
                pavimento = ""
                if 'pavimento' in texto.lower() or 'térreo' in texto.lower():
                    pavimento = extrair_pavimento_otimizado(texto)
                
                dados.append({
                    "Formato": "Bloco",
                    "Apartamento": numero,
                    "Tipo": tipo,
                    "Torre/Bloco": bloco,
                    "Pavimento": pavimento,
                    "Área Privativa (m²)": privativa.replace(",", "."),
                    "Área Comum (m²)": comum.replace(",", "."),
                    "Área Total (m²)": total.replace(",", "."),
                    "Fração Ideal (%)": fracao.replace(",", "."),
                    "Área Terreno (m²)": equivalente.replace(",", "."),
                    "Descrição": descricao
                })
    
    return pd.DataFrame(dados)

def extrair_casas_otimizado(doc):
    """Versão ULTRA-OTIMIZADA da extração de casas - MÁXIMA PERFORMANCE"""
    # Concatenar texto de forma otimizada
    textos = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto and len(texto) > 10 and ('CASA' in texto or 'Casa' in texto):  # Filtro mais agressivo
            textos.append(texto)
    
    texto_completo = "\n".join(textos)
    blocos = REGEX_CASA_SPLIT.split(texto_completo)
    casas = [(blocos[i], blocos[i + 1]) for i in range(1, len(blocos), 2)]

    dados = []
    for numero, conteudo in casas:
        if not conteudo or len(conteudo) < 30:  # Pular conteúdo muito pequeno
            continue
            
        # Executar todos os regex de uma vez para melhor performance
        area_terreno = REGEX_AREA_TERRENO.search(conteudo)
        area_construida = REGEX_AREA_CONSTRUIDA.search(conteudo)
        area_comum = REGEX_AREA_COMUM.search(conteudo)
        area_total = REGEX_AREA_TOTAL_REAL.search(conteudo)
        fracao = REGEX_FRACAO_IDEAL.search(conteudo)
        
        # Extrair descrição de forma otimizada - apenas se necessário
        descricao = ""
        if 'frente' in conteudo.lower():
            idx_frente = conteudo.lower().find('frente')
            idx_ultimo_ponto = conteudo.rfind('.')
            if idx_ultimo_ponto > idx_frente:
                descricao = conteudo[idx_frente:idx_ultimo_ponto+1].replace('\n', ' ').strip()
        
        # Extrair pavimento de forma otimizada - apenas se necessário
        pavimento = ""
        if 'pavimento' in conteudo.lower() or 'térreo' in conteudo.lower():
            pavimento = extrair_pavimento_otimizado(conteudo)
        
        dados.append({
            "Formato": "Casa",
            "Número da Casa": numero,
            "Pavimento": pavimento,
            "Área do Terreno (m²)": area_terreno.group(1) if area_terreno else "",
            "Área Construída (m²)": area_construida.group(1) if area_construida else "",
            "Área Comum Real (m²)": area_comum.group(1) if area_comum else "",
            "Área Total Real (m²)": area_total.group(1) if area_total else "",
            "Fração Ideal (%)": (fracao.group(1) + " %") if fracao else "",
            "Descrição": descricao
        })

    return pd.DataFrame(dados)

def processar_arquivo_otimizado(path):
    """Versão ULTRA-OTIMIZADA do processamento de arquivo - MÁXIMA PERFORMANCE"""
    try:
        # Verificar se é PDF ou DOCX
        if path.lower().endswith('.pdf'):
            texto_completo = extrair_texto_pdf(path)
            if not texto_completo.strip():
                return None
            
            # Criar objeto fake para compatibilidade - FILTRO MAIS AGRESSIVO
            linhas_validas = []
            for t in texto_completo.split('\n'):
                t = t.strip()
                if t and len(t) > 10 and ('Apartamento' in t or 'APARTAMENTO' in t or 'CASA' in t or 'Casa' in t):
                    linhas_validas.append(t)
            
            class FakeDoc:
                def __init__(self, linhas):
                    self.paragraphs = [type('p', (), {'text': t}) for t in linhas]
            doc = FakeDoc(linhas_validas)
        else:
            # Arquivo DOCX com otimização ULTRA de memória
            doc = Document(path)
            # Filtro ULTRA agressivo - apenas parágrafos relevantes
            paragrafos_validos = []
            for p in doc.paragraphs:
                texto = p.text.strip()
                if texto and len(texto) > 15 and ('Apartamento' in texto or 'APARTAMENTO' in texto or 'CASA' in texto or 'Casa' in texto):
                    paragrafos_validos.append(texto)
            
            texto_completo = "\n".join(paragrafos_validos)
            
            # Criar objeto otimizado
            class OptimizedDoc:
                def __init__(self, paragrafos):
                    self.paragraphs = [type('p', (), {'text': t}) for t in paragrafos]
            doc = OptimizedDoc(paragrafos_validos)
        
        # Identificar tipo de documento
        tipo = identificar_tipo_documento_otimizado(texto_completo)
        
        # Processar baseado no tipo
        if tipo == "torre":
            df = extrair_torre_otimizado(doc)
        elif tipo == "bloco":
            df = extrair_bloco_otimizado(doc)
        elif tipo == "casa":
            df = extrair_casas_otimizado(doc)
        else:
            return None
        
        if df is not None and not df.empty:
            return df
        else:
            return None
    except Exception as e:
        print(f"Erro ao processar arquivo {path}: {e}")
        return None

def processar_arquivos_paralelo(arquivos, max_workers=None):
    """Processa múltiplos arquivos em paralelo, otimizando uso de memória."""
    todos_dfs = []
    documentos_processados = []
    if max_workers is None:
        try:
            import multiprocessing
            # Aumentar workers para melhor performance
            max_workers = min(8, multiprocessing.cpu_count())
        except Exception:
            max_workers = 4
    
    # Para arquivos únicos, usar processamento sequencial mais rápido
    if len(arquivos) == 1:
        df = processar_arquivo_otimizado(arquivos[0])
        if df is not None and not df.empty:
            documentos_processados.append({
                'filename': os.path.basename(arquivos[0]),
                'rows_extracted': len(df),
                'tipo_documento': df['Formato'].iloc[0] if 'Formato' in df.columns else 'desconhecido'
            })
            return df, documentos_processados
        else:
            documentos_processados.append({
                'filename': os.path.basename(arquivos[0]),
                'rows_extracted': 0,
                'error': 'Nenhum dado extraído'
            })
            return pd.DataFrame(), documentos_processados
    
    # Para múltiplos arquivos, usar processamento paralelo
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        future_to_arquivo = {executor.submit(processar_arquivo_otimizado, arq): arq for arq in arquivos}
        for future in as_completed(future_to_arquivo):
            arquivo = future_to_arquivo[future]
            try:
                df = future.result()
                if df is not None and not df.empty:
                    todos_dfs.append(df)
                    documentos_processados.append({
                        'filename': os.path.basename(arquivo),
                        'rows_extracted': len(df),
                        'tipo_documento': df['Formato'].iloc[0] if 'Formato' in df.columns else 'desconhecido'
                    })
                else:
                    documentos_processados.append({
                        'filename': os.path.basename(arquivo),
                        'rows_extracted': 0,
                        'error': 'Nenhum dado extraído'
                    })
            except Exception as e:
                documentos_processados.append({
                    'filename': os.path.basename(arquivo),
                    'rows_extracted': 0,
                    'error': str(e)
                })
    if todos_dfs:
        resultado = pd.concat(todos_dfs, ignore_index=True)
        return resultado, documentos_processados
    else:
        return pd.DataFrame(), documentos_processados

def main():
    parser = argparse.ArgumentParser(description="Extrai dados de memorial de incorporação de vários arquivos DOCX/PDF (versão ultra-otimizada).")
    parser.add_argument('arquivos', nargs='+', help='Caminhos dos arquivos .docx/.pdf ou diretório contendo arquivos')
    parser.add_argument('-o', '--output', default='dados_memorial_otimizado.xlsx', help='Nome do arquivo Excel de saída')
    parser.add_argument('-w', '--workers', type=int, default=4, help='Número de workers para processamento paralelo')
    args = parser.parse_args()

    # Coletar arquivos
    arquivos = []
    for entrada in args.arquivos:
        if os.path.isdir(entrada):
            arquivos.extend([os.path.join(entrada, f) for f in os.listdir(entrada) 
                           if f.lower().endswith(('.docx', '.pdf'))])
        else:
            arquivos.append(entrada)
    
    if not arquivos:
        print("Nenhum arquivo .docx ou .pdf encontrado.")
        sys.exit(1)

    print(f"🚀 Processando {len(arquivos)} arquivos com {args.workers} workers...")
    start_time = time.time()
    
    # Processar arquivos em paralelo
    resultado, documentos_processados = processar_arquivos_paralelo(arquivos, args.workers)
    
    if not resultado.empty:
        resultado.to_excel(args.output, index=False)
    else:
        print("Nenhum dado extraído dos arquivos informados.")
        sys.exit(1)
    
    end_time = time.time()
    processing_time = end_time - start_time
    
    print(f"\n✅ Extração concluída em {processing_time:.2f}s!")
    print(f"📊 Total de registros: {len(resultado)}")
    print(f"📁 Dados salvos em: {args.output}")

if __name__ == "__main__":
    main() 