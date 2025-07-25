# Endpoints relacionados ao ChatGPT/OpenAI

from flask import Blueprint, request, jsonify, send_file, current_app
import os
import uuid
import shutil
import re
import tempfile
from werkzeug.utils import secure_filename
from ai.openai_service import extract_fields_with_openai
import pypdf
from config import Config
from security import secure_manager
import pandas as pd

ai_bp = Blueprint('ai', __name__)

# Dicionário global para armazenar arquivos de memorial
memorial_files = {}

def allowed_file(filename):
    """Verifica se o arquivo tem extensão permitida"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_EXTENSIONS

@ai_bp.route('/api/process-file', methods=['POST'])
def process_file_chatgpt():
    """Endpoint otimizado para processamento com ChatGPT - agora SEMPRE faz OCR antes da IA"""
    from ai.ocr_service import extract_text_from_pdf, process_pdf_with_ocr
    temp_file_path = None
    user_ip = request.remote_addr
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Apenas arquivos PDF são permitidos'}), 400
        # Obter tipo de serviço
        service_type = request.form.get('service', 'matricula')
        print(f"🎯 Serviço recebido: {service_type}")
        original_filename = secure_filename(file.filename or 'unknown.pdf')
        # Processar arquivo de forma segura
        if Config.SECURE_PROCESSING:
            temp_file_path, file_id = secure_manager.process_file_securely(
                file, original_filename, user_ip
            )
        else:
            file_id = str(uuid.uuid4())
            upload_filename = f"{file_id}_{original_filename}"
            temp_file_path = os.path.join(Config.UPLOAD_FOLDER, upload_filename)
            file.save(temp_file_path)
        # DESCRIPTOGRAFAR ANTES DO OCR!
        if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
            temp_file_path = secure_manager.decrypt_file(temp_file_path)
        # Sempre rodar OCR antes da IA
        text_content = ""
        temp_ocr_path = temp_file_path + '_ocr.pdf'
        try:
            ocr_result = process_pdf_with_ocr(temp_file_path, temp_ocr_path)
            if ocr_result.get('success'):
                print("✅ OCR bem-sucedido, extraindo texto...")
                text_content = extract_text_from_pdf(temp_ocr_path)
                text_content = re.sub(r'\s+', ' ', text_content.replace('\n', ' ')).strip()
            else:
                print(f"❌ OCR falhou: {ocr_result.get('error', 'Erro desconhecido')}")
                return jsonify({
                    'error': 'Não foi possível extrair texto suficiente do PDF.',
                    'details': 'O arquivo pode estar corrompido, protegido por senha, ou ser uma imagem escaneada de baixa qualidade.',
                    'suggestion': 'Tente com um arquivo PDF diferente ou verifique se o arquivo não está protegido.',
                    'ocr_error': ocr_result.get('error', 'Erro desconhecido')
                }), 400
        except Exception as ocr_error:
            print(f"❌ Erro durante OCR: {str(ocr_error)}")
            return jsonify({
                'error': 'Erro durante processamento OCR.',
                'details': str(ocr_error),
                'suggestion': 'Verifique se o arquivo é um PDF válido e não está corrompido.'
            }), 400
        # Verificar se temos texto suficiente para processar
        if not text_content or len(text_content.strip()) < 10:
            return jsonify({
                'error': 'Texto insuficiente para processamento.',
                'details': f'Extraído apenas {len(text_content)} caracteres.',
                'suggestion': 'O arquivo pode estar vazio ou não conter texto legível.',
                'text_preview': text_content[:200] if text_content else ''
            }), 400
        print(f"✅ Texto extraído com sucesso: {len(text_content)} caracteres")
        print(f"📄 Preview do texto (primeiros 500 chars): {text_content[:500]}")
        # Extrair campos com OpenAI
        print("🤖 Iniciando extração com IA...")
        model = request.form.get('model', 'gpt-4o')
        print(f"🎯 Modelo recebido no endpoint /api/process-file: {model}")
        campos = extract_fields_with_openai(text_content, model=model, service_type=service_type)
        # Limpar arquivo temporário após processamento
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        if 'temp_ocr_path' in locals() and temp_ocr_path and os.path.exists(temp_ocr_path):
            try:
                os.remove(temp_ocr_path)
            except Exception:
                pass
        return jsonify({
            'success': True,
            'message': f'PDF processado e campos extraídos com ChatGPT ({service_type})!',
            'original_filename': original_filename,
            'file_id': file_id,
            'campos': campos,
            'model': model,
            'service_type': service_type,
            'text_length': len(text_content),
            'used_ocr_fallback': True,
            'secure_processing': Config.SECURE_PROCESSING
        })
    except Exception as e:
        # Garantir limpeza em caso de erro
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        return jsonify({'error': f'Erro interno do servidor: {str(e)}'}), 500

@ai_bp.route('/api/certidao', methods=['POST'])
def process_certidao():
    """Endpoint para processar PDF de matrícula, extrair campos via OpenAI e gerar certidão personalizada"""
    from ai.ocr_service import extract_text_from_pdf, process_pdf_with_ocr
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    import io

    temp_file_path = None
    user_ip = request.remote_addr

    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Apenas arquivos PDF são permitidos'}), 400

        # Obter tipo de certidão e modelo
        tipo_certidao = 'STNegativa'
        motivo_certidao = 'Não há ônus real identificado.'

        original_filename = secure_filename(file.filename or 'unknown.pdf')
        # Processar arquivo de forma segura
        if Config.SECURE_PROCESSING:
            temp_file_path, file_id = secure_manager.process_file_securely(
                file, original_filename, user_ip
            )
        else:
            file_id = str(uuid.uuid4())
            upload_filename = f"{file_id}_{original_filename}"
            temp_file_path = os.path.join(Config.UPLOAD_FOLDER, upload_filename)
            file.save(temp_file_path)

        # DESCRIPTOGRAFAR ANTES DO OCR!
        if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
            temp_file_path = secure_manager.decrypt_file(temp_file_path)

        # Sempre rodar OCR antes da IA
        text_content = ""
        temp_ocr_path = temp_file_path + '_ocr.pdf'
        try:
            ocr_result = process_pdf_with_ocr(temp_file_path, temp_ocr_path)
            if ocr_result.get('success'):
                print("✅ OCR bem-sucedido, extraindo texto...")
                text_content = extract_text_from_pdf(temp_ocr_path)
                text_content = re.sub(r'\s+', ' ', text_content.replace('\n', ' ')).strip()
            else:
                print(f"❌ OCR falhou: {ocr_result.get('error', 'Erro desconhecido')}")
                return jsonify({
                    'error': 'Não foi possível extrair texto suficiente do PDF.',
                    'details': 'O arquivo pode estar corrompido, protegido por senha, ou ser uma imagem escaneada de baixa qualidade.',
                    'suggestion': 'Tente com um arquivo PDF diferente ou verifique se o arquivo não está protegido.',
                    'ocr_error': ocr_result.get('error', 'Erro desconhecido')
                }), 400
        except Exception as ocr_error:
            print(f"❌ Erro durante OCR: {str(ocr_error)}")
            return jsonify({
                'error': 'Erro durante processamento OCR.',
                'details': str(ocr_error),
                'suggestion': 'Verifique se o arquivo é um PDF válido e não está corrompido.'
            }), 400

        # Verificar se temos texto suficiente para processar
        if not text_content or len(text_content.strip()) < 10:
            return jsonify({
                'error': 'Texto insuficiente para processamento.',
                'details': f'Extraído apenas {len(text_content)} caracteres.',
                'suggestion': 'O arquivo pode estar vazio ou não conter texto legível.',
                'text_preview': text_content[:200] if text_content else ''
            }), 400

        print(f"✅ Texto extraído com sucesso: {len(text_content)} caracteres")
        print(f"📄 Preview do texto (primeiros 500 chars): {text_content[:500]}")

        # Extrair campos com OpenAI
        print("🤖 Iniciando extração com IA...")
        model = request.form.get('model', 'gpt-4o')
        print(f"🎯 Modelo recebido no endpoint /api/certidao (PDF): {model}")
        if not Config.OPENAI_API_KEY or Config.OPENAI_API_KEY.strip() == '':
            return jsonify({
                'error': 'Chave da API OpenAI não configurada.',
                'details': 'Configure a variável de ambiente OPENAI_API_KEY para usar a funcionalidade de IA.',
                'suggestion': 'Adicione sua chave da OpenAI nas configurações do sistema.'
            }), 500

        campos = None
        try:
            campos = extract_fields_with_openai(text_content, model=model, service_type='certidao')
        except Exception as ia_error:
            print(f"❌ Erro na extração de campos pela IA: {ia_error}")
            return jsonify({'error': 'Erro ao extrair campos da certidão com IA.', 'details': str(ia_error)}), 500

        # --- Lógica para identificar o tipo de certidão ---
        if not campos or not isinstance(campos, dict):
            return jsonify({'error': 'Não foi possível extrair os campos necessários da certidão.'}), 500
        tipo_certidao = 'STNegativa'
        motivo_certidao = 'Não há ônus real identificado.'
        onus = (campos.get('onus_certidao_negativa') or '').lower()
        descricao = (campos.get('descricao_imovel') or '').lower()
        enfiteuta = (campos.get('enfiteuta') or '').lower()
        senhorio = (campos.get('senhorio_direto') or '').lower()
        if 'foreiro' in descricao or 'enfiteuta' in descricao or enfiteuta or senhorio:
            tipo_certidao = 'STForeiro'
            motivo_certidao = 'Imóvel foreiro (domínio útil/enfiteuta) identificado.'
        elif any(palavra in onus for palavra in ['hipoteca', 'alienação', 'penhora', 'ônus', 'restrição', 'gravame', 'fiduciária', 'ação judicial', 'usucapião', 'usufruto', 'servidão', 'penhor', 'protesto', 'bloqueio', 'penalidade', 'penal', 'ação', 'execução']):
            tipo_certidao = 'STPositiva'
            motivo_certidao = 'Ônus real ou restrição identificado.'
        else:
            tipo_certidao = 'STNegativa'
            motivo_certidao = 'Não há ônus real identificado.'

        # O restante do código permanece igual, mas usar tipo_certidao para o título
        # e retornar tipo_certidao e motivo_certidao no response.

        # Gerar PDF personalizado da certidão conforme o tipo (layout avançado)
        from reportlab.lib.units import cm
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase import pdfmetrics
        from datetime import datetime
        from reportlab.lib import colors
        import io

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
            leftMargin=3*cm, rightMargin=3*cm, topMargin=3*cm, bottomMargin=2*cm)

        # Fonte Times New Roman se disponível
        try:
            pdfmetrics.registerFont(TTFont('TimesNewRoman', 'times.ttf'))
            base_font = 'TimesNewRoman'
        except:
            base_font = 'Times-Roman'

        # Estilos
        style_title = ParagraphStyle('title', fontName=base_font, fontSize=13, alignment=TA_CENTER, leading=18, spaceAfter=10, textDecoration='underline', fontWeight='bold')
        style_body = ParagraphStyle('body', fontName=base_font, fontSize=12, alignment=TA_JUSTIFY, leading=19)

        # Cores
        cor_vermelho = '#FF0000'
        cor_azul = '#0070C0'
        cor_azul_escuro = '#002060'

        # Data
        data_certidao = datetime.now().strftime('%d/%m/%Y')

        # Montar o parágrafo principal
        texto = (
            '<b>CERTIFICO</b>, nos termos dos arts. 17 e 19, §9º, da Lei n.º 6.015/1973, e art. 123, caput, do Provimento n.º 149/2023, do Conselho Nacional de Justiça - CNJ, que, revendo os livros, arquivos e sistemas eletrônicos desta Serventia, inclusive cadastro interno de ações reais e pessoais reipersecutórias envolvendo imóveis desta circunscrição, encontrei o lançamento relativo ao registro de imóvel seguinte: '
            f'<b><u>CADASTRO NACIONAL DE MATRÍCULA - CNM:</u></b> <font color="{cor_vermelho}">{campos.get("cnm", "")}</font>, '
            f'<b><u>DESCRIÇÃO DO IMÓVEL:</u></b> <font color="{cor_azul}">{campos.get("descricao_imovel", "")}</font>, '
        )
        if campos.get('senhorio_direto'):
            texto += f'<b><u>SENHORIO DIRETO:</u></b> {campos.get("senhorio_direto")}, '
        if campos.get('enfiteuta'):
            texto += f'<b><u>ENFITEUTA:</u></b> {campos.get("enfiteuta")}, '
        texto += f'<b><u>PROPRIETÁRIO(S):</u></b> <font color="{cor_azul_escuro}">{campos.get("proprietarios", "")}</font>, '
        texto += f'<b><u>INSCRIÇÃO IMOBILIÁRIA:</u></b> <font color="{cor_azul_escuro}">{campos.get("inscricao_imobiliaria", "")}</font>, '
        if campos.get('rip'):
            texto += f'<b><u>REGISTRO IMOBILIÁRIO PATRIMONIAL (RIP):</u></b> {campos.get("rip")}, '
        texto += f'<b><u>DIREITOS, ÔNUS REAIS E RESTRIÇÕES JUDICIAIS E ADMINISTRATIVAS:</u></b> <b>{campos.get("onus_certidao_negativa", "")}</b>. '
        texto += f'O referido é verdade e dou fé. São Luís/MA, {data_certidao}. '
        texto += f'<b>Emolumentos:</b> Certidão: Ato 16.24.4 - <font color="{cor_azul}">R$ 87,31</font>; FEMP: <font color="{cor_azul}">R$ 3,49</font>; FADEP: <font color="{cor_azul}">R$ 3,49</font>; FERC: <font color="{cor_azul}">R$ 2,61</font>. '
        texto += '<b>João Gabriel Santos Barros, Escrevente Autorizado.</b> '
        texto += f'<br/><b><font color="{cor_azul_escuro}">Validade: 30 dias.</font></b>'

        # Ajustar tags HTML para ReportLab
        texto = texto.replace('class="bold"', 'b').replace('class="underline"', 'u')
        texto = texto.replace('b u', 'b u').replace('u b', 'u b')  # garantir ordem
        texto = texto.replace('style="color:'+cor_azul_escuro+'"', f'color="{cor_azul_escuro}"')

        from reportlab.platypus import Flowable
        
        elements: list[Flowable] = [
            Paragraph('<u><b>CERTIDÃO DE SITUAÇÃO JURÍDICA DO IMÓVEL</b></u>', style_title),
            Paragraph(texto, style_body)
        ]

        doc.build(elements)
        buffer.seek(0)

        # Limpar arquivos temporários
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        if 'temp_ocr_path' in locals() and temp_ocr_path and os.path.exists(temp_ocr_path):
            try:
                os.remove(temp_ocr_path)
            except Exception:
                pass

        # Debug: mostrar campos extraídos antes de gerar o PDF
        print("Campos extraídos para o PDF:", campos)
        # Mapeamento para garantir nomes corretos
        campos['cnm'] = campos.get('cnm') or campos.get('matricula') or ''
        campos['descricao_imovel'] = campos.get('descricao_imovel') or campos.get('descricao') or ''
        campos['proprietarios'] = campos.get('proprietarios') or campos.get('proprietario') or ''
        campos['inscricao_imobiliaria'] = campos.get('inscricao_imobiliaria') or campos.get('inscricao') or ''
        campos['onus_certidao_negativa'] = campos.get('onus_certidao_negativa') or campos.get('onus') or ''
        campos['senhorio_enfiteuta'] = campos.get('senhorio_enfiteuta') or campos.get('senhorio') or campos.get('enfiteuta') or ''
        campos['rip'] = campos.get('rip') or ''
        campos['nome_solicitante'] = campos.get('nome_solicitante') or campos.get('solicitante') or ''

        # Retornar PDF gerado e info do tipo/motivo
        import json
        campos_json = json.dumps(campos)
        
        return (buffer.getvalue(), 200, {
            'Content-Type': 'application/pdf',
            'Content-Disposition': f'attachment; filename=certidao_{tipo_certidao}_{file_id}.pdf',
            'X-Certidao-Tipo': tipo_certidao,
            'X-Certidao-Motivo': motivo_certidao,
            'X-Certidao-Data': campos_json  # Incluir dados extraídos nos headers
        })
        
    except Exception as e:
        print(f"❌ Erro geral na geração da certidão: {str(e)}")
        # Limpar arquivos temporários em caso de erro
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        if 'temp_ocr_path' in locals() and temp_ocr_path and os.path.exists(temp_ocr_path):
            try:
                os.remove(temp_ocr_path)
            except Exception:
                pass
        return jsonify({'error': f'Erro interno do servidor: {str(e)}'}), 500

@ai_bp.route('/api/certidao/data', methods=['POST'])
def get_certidao_data():
    """Endpoint para extrair dados da certidão para download Word"""
    from ai.ocr_service import extract_text_from_pdf, process_pdf_with_ocr
    from ai.openai_service import extract_fields_with_openai
    import io

    temp_file_path = None
    user_ip = request.remote_addr

    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Apenas arquivos PDF são permitidos'}), 400

        original_filename = secure_filename(file.filename or 'unknown.pdf')
        # Processar arquivo de forma segura
        if Config.SECURE_PROCESSING:
            temp_file_path, file_id = secure_manager.process_file_securely(
                file, original_filename, user_ip
            )
        else:
            file_id = str(uuid.uuid4())
            upload_filename = f"{file_id}_{original_filename}"
            temp_file_path = os.path.join(Config.UPLOAD_FOLDER, upload_filename)
            file.save(temp_file_path)

        # Extrair texto do PDF (OCR se necessário)
        text_content = ""
        temp_ocr_path = None
        
        try:
            # Descriptografar se necessário
            if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
                temp_file_path = secure_manager.decrypt_file(temp_file_path)
            
            # Tentar extrair texto diretamente primeiro
            with open(temp_file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
        except Exception as e:
            print(f"❌ Erro ao extrair texto diretamente: {str(e)}")
            text_content = ""
        
        # Limpar e normalizar texto
        text_content = re.sub(r'\s+', ' ', text_content.replace('\n', ' ')).strip()
        
        # Se não houver texto suficiente, tentar OCR
        if not text_content or len(text_content.strip()) < 50:
            print("🔍 Texto insuficiente, tentando OCR...")
            try:
                if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
                    temp_file_path = secure_manager.decrypt_file(temp_file_path)
                
                text_content = process_pdf_with_ocr(temp_file_path, user_ip)
                if not text_content:
                    return jsonify({'error': 'Não foi possível extrair texto do PDF'}), 400
            except Exception as ocr_error:
                print(f"❌ Erro no OCR: {str(ocr_error)}")
                return jsonify({'error': 'Erro no processamento OCR'}), 500

        # Extrair campos usando OpenAI
        print("🔍 Extraindo campos da certidão...")
        model = request.form.get('model', 'gpt-4o')
        print(f"🎯 Modelo recebido no endpoint /api/certidao/data: {model}")
        campos = extract_fields_with_openai(text_content, model=model, service_type="certidao")
        
        if not campos or 'error' in campos:
            error_msg = campos.get('error', 'Erro desconhecido na extração') if campos else 'Nenhum dado extraído'
            return jsonify({'error': f'Erro na extração de dados: {error_msg}'}), 500

        # Limpar arquivos temporários
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        if 'temp_ocr_path' in locals() and temp_ocr_path and os.path.exists(temp_ocr_path):
            try:
                os.remove(temp_ocr_path)
            except Exception:
                pass

        return jsonify({
            'success': True,
            'data': campos,
            'message': 'Dados da certidão extraídos com sucesso'
        })

    except Exception as e:
        print(f"❌ Erro geral na extração de dados da certidão: {str(e)}")
        # Limpar arquivos temporários em caso de erro
        if Config.SECURE_PROCESSING and temp_file_path:
            secure_manager.cleanup_file(temp_file_path, user_ip)
        if 'temp_ocr_path' in locals() and temp_ocr_path and os.path.exists(temp_ocr_path):
            try:
                os.remove(temp_ocr_path)
            except Exception:
                pass
        return jsonify({'error': f'Erro interno: {str(e)}'}), 500 

@ai_bp.route('/api/certidao/word', methods=['POST'])
def generate_certidao_word():
    """Endpoint para gerar arquivo Word da certidão com exatamente a mesma formatação do PDF"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        import io
        from datetime import datetime
        
        # Obter dados da requisição
        data = request.get_json()
        if not data or 'data' not in data:
            return jsonify({'error': 'Dados da certidão não fornecidos'}), 400
        
        certidao_data = data['data']
        
        # Criar documento Word
        doc = Document()
        
        # Configurar margens (exatamente iguais ao PDF)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.2)      # 3cm
            section.bottom_margin = Inches(0.8)   # 2cm
            section.left_margin = Inches(1.2)     # 3cm
            section.right_margin = Inches(1.2)    # 3cm
        
        # Data
        data_certidao = datetime.now().strftime('%d/%m/%Y')
        
        # Cores (exatamente iguais ao PDF)
        cor_vermelho = RGBColor(255, 0, 0)      # #FF0000
        cor_azul = RGBColor(0, 112, 192)        # #0070C0
        cor_azul_escuro = RGBColor(0, 32, 96)   # #002060
        
        # Adicionar título centralizado e sublinhado
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run("CERTIDÃO DE SITUAÇÃO JURÍDICA DO IMÓVEL")
        title_run.bold = True
        title_run.underline = True
        title_run.font.size = Pt(13)
        title_run.font.name = 'Times New Roman'
        
        # Adicionar espaço após título
        doc.add_paragraph()
        
        # Adicionar parágrafo principal justificado
        main_paragraph = doc.add_paragraph()
        main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Adicionar texto com formatação exata do PDF
        # Início do texto
        run1 = main_paragraph.add_run("CERTIFICO, nos termos dos arts. 17 e 19, §9º, da Lei n.º 6.015/1973, e art. 123, caput, do Provimento n.º 149/2023, do Conselho Nacional de Justiça - CNJ, que, revendo os livros, arquivos e sistemas eletrônicos desta Serventia, inclusive cadastro interno de ações reais e pessoais reipersecutórias envolvendo imóveis desta circunscrição, encontrei o lançamento relativo ao registro de imóvel seguinte: ")
        run1.font.size = Pt(12)
        run1.font.name = 'Times New Roman'
        
        # CNM em vermelho
        run2 = main_paragraph.add_run("CADASTRO NACIONAL DE MATRÍCULA - CNM: ")
        run2.bold = True
        run2.underline = True
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        
        run3 = main_paragraph.add_run(f"{certidao_data.get('cnm', '')}")
        run3.font.color.rgb = cor_vermelho
        run3.font.size = Pt(12)
        run3.font.name = 'Times New Roman'
        
        run4 = main_paragraph.add_run(", ")
        run4.font.size = Pt(12)
        run4.font.name = 'Times New Roman'
        
        # Descrição do imóvel em azul
        run5 = main_paragraph.add_run("DESCRIÇÃO DO IMÓVEL: ")
        run5.bold = True
        run5.underline = True
        run5.font.size = Pt(12)
        run5.font.name = 'Times New Roman'
        
        run6 = main_paragraph.add_run(f"{certidao_data.get('descricao_imovel', '')}")
        run6.font.color.rgb = cor_azul
        run6.font.size = Pt(12)
        run6.font.name = 'Times New Roman'
        
        run7 = main_paragraph.add_run(", ")
        run7.font.size = Pt(12)
        run7.font.name = 'Times New Roman'
        
        # Senhorio direto (se existir)
        if certidao_data.get('senhorio_direto'):
            run8 = main_paragraph.add_run("SENHORIO DIRETO: ")
            run8.bold = True
            run8.underline = True
            run8.font.size = Pt(12)
            run8.font.name = 'Times New Roman'
            
            run9 = main_paragraph.add_run(f"{certidao_data.get('senhorio_direto')}")
            run9.font.size = Pt(12)
            run9.font.name = 'Times New Roman'
            
            run10 = main_paragraph.add_run(", ")
            run10.font.size = Pt(12)
            run10.font.name = 'Times New Roman'
        
        # Enfiteuta (se existir)
        if certidao_data.get('enfiteuta'):
            run11 = main_paragraph.add_run("ENFITEUTA: ")
            run11.bold = True
            run11.underline = True
            run11.font.size = Pt(12)
            run11.font.name = 'Times New Roman'
            
            run12 = main_paragraph.add_run(f"{certidao_data.get('enfiteuta')}")
            run12.font.size = Pt(12)
            run12.font.name = 'Times New Roman'
            
            run13 = main_paragraph.add_run(", ")
            run13.font.size = Pt(12)
            run13.font.name = 'Times New Roman'
        
        # Proprietários em azul escuro
        run14 = main_paragraph.add_run("PROPRIETÁRIO(S): ")
        run14.bold = True
        run14.underline = True
        run14.font.size = Pt(12)
        run14.font.name = 'Times New Roman'
        
        run15 = main_paragraph.add_run(f"{certidao_data.get('proprietarios', '')}")
        run15.font.color.rgb = cor_azul_escuro
        run15.font.size = Pt(12)
        run15.font.name = 'Times New Roman'
        
        run16 = main_paragraph.add_run(", ")
        run16.font.size = Pt(12)
        run16.font.name = 'Times New Roman'
        
        # Inscrição imobiliária em azul escuro
        run17 = main_paragraph.add_run("INSCRIÇÃO IMOBILIÁRIA: ")
        run17.bold = True
        run17.underline = True
        run17.font.size = Pt(12)
        run17.font.name = 'Times New Roman'
        
        run18 = main_paragraph.add_run(f"{certidao_data.get('inscricao_imobiliaria', '')}")
        run18.font.color.rgb = cor_azul_escuro
        run18.font.size = Pt(12)
        run18.font.name = 'Times New Roman'
        
        run19 = main_paragraph.add_run(", ")
        run19.font.size = Pt(12)
        run19.font.name = 'Times New Roman'
        
        # RIP (se existir)
        if certidao_data.get('rip'):
            run20 = main_paragraph.add_run("REGISTRO IMOBILIÁRIO PATRIMONIAL (RIP): ")
            run20.bold = True
            run20.underline = True
            run20.font.size = Pt(12)
            run20.font.name = 'Times New Roman'
            
            run21 = main_paragraph.add_run(f"{certidao_data.get('rip')}")
            run21.font.size = Pt(12)
            run21.font.name = 'Times New Roman'
            
            run22 = main_paragraph.add_run(", ")
            run22.font.size = Pt(12)
            run22.font.name = 'Times New Roman'
        
        # Ônus em negrito
        run23 = main_paragraph.add_run("DIREITOS, ÔNUS REAIS E RESTRIÇÕES JUDICIAIS E ADMINISTRATIVAS: ")
        run23.bold = True
        run23.underline = True
        run23.font.size = Pt(12)
        run23.font.name = 'Times New Roman'
        
        run24 = main_paragraph.add_run(f"{certidao_data.get('onus_certidao_negativa', '')}")
        run24.bold = True
        run24.font.size = Pt(12)
        run24.font.name = 'Times New Roman'
        
        run25 = main_paragraph.add_run(". ")
        run25.font.size = Pt(12)
        run25.font.name = 'Times New Roman'
        
        # Finalização
        run26 = main_paragraph.add_run(f"O referido é verdade e dou fé. São Luís/MA, {data_certidao}. ")
        run26.font.size = Pt(12)
        run26.font.name = 'Times New Roman'
        
        # Emolumentos
        run27 = main_paragraph.add_run("Emolumentos: ")
        run27.bold = True
        run27.font.size = Pt(12)
        run27.font.name = 'Times New Roman'
        
        run28 = main_paragraph.add_run("Certidão: Ato 16.24.4 - ")
        run28.font.size = Pt(12)
        run28.font.name = 'Times New Roman'
        
        run29 = main_paragraph.add_run("R$ 87,31")
        run29.font.color.rgb = cor_azul
        run29.font.size = Pt(12)
        run29.font.name = 'Times New Roman'
        
        run30 = main_paragraph.add_run("; FEMP: ")
        run30.font.size = Pt(12)
        run30.font.name = 'Times New Roman'
        
        run31 = main_paragraph.add_run("R$ 3,49")
        run31.font.color.rgb = cor_azul
        run31.font.size = Pt(12)
        run31.font.name = 'Times New Roman'
        
        run32 = main_paragraph.add_run("; FADEP: ")
        run32.font.size = Pt(12)
        run32.font.name = 'Times New Roman'
        
        run33 = main_paragraph.add_run("R$ 3,49")
        run33.font.color.rgb = cor_azul
        run33.font.size = Pt(12)
        run33.font.name = 'Times New Roman'
        
        run34 = main_paragraph.add_run("; FERC: ")
        run34.font.size = Pt(12)
        run34.font.name = 'Times New Roman'
        
        run35 = main_paragraph.add_run("R$ 2,61")
        run35.font.color.rgb = cor_azul
        run35.font.size = Pt(12)
        run35.font.name = 'Times New Roman'
        
        run36 = main_paragraph.add_run(". ")
        run36.font.size = Pt(12)
        run36.font.name = 'Times New Roman'
        
        # Assinatura
        run37 = main_paragraph.add_run("João Gabriel Santos Barros, Escrevente Autorizado.")
        run37.bold = True
        run37.font.size = Pt(12)
        run37.font.name = 'Times New Roman'
        
        # Adicionar quebra de linha
        main_paragraph.add_run("\n")
        
        # Validade em azul escuro
        run38 = main_paragraph.add_run("Validade: 30 dias.")
        run38.bold = True
        run38.font.color.rgb = cor_azul_escuro
        run38.font.size = Pt(12)
        run38.font.name = 'Times New Roman'
        
        # Salvar documento em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'certidao_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Erro ao gerar arquivo Word: {str(e)}")
        return jsonify({'error': f'Erro ao gerar arquivo Word: {str(e)}'}), 500

@ai_bp.route('/api/qualificacao', methods=['POST'])
def process_qualificacao():
    """Endpoint para processar múltiplos arquivos para análise de qualificação"""
    import uuid
    request_id = str(uuid.uuid4())[:8]
    print(f"🆔 NOVA REQUISIÇÃO /api/qualificacao - ID: {request_id}")
    print(f"📊 Headers: {dict(request.headers)}")
    print(f"🌐 User-Agent: {request.headers.get('User-Agent', 'N/A')}")
    
    temp_files = []
    user_ip = request.remote_addr
    
    try:
        if 'files[]' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        
        files = request.files.getlist('files[]')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        
        # Validar arquivos
        valid_files = []
        for file in files:
            if file.filename and allowed_file(file.filename):
                valid_files.append(file)
            else:
                return jsonify({'error': f'Arquivo inválido: {file.filename}. Apenas PDFs são permitidos.'}), 400
        
        if not valid_files:
            return jsonify({'error': 'Nenhum arquivo válido encontrado'}), 400
        
        print(f"📁 Processando {len(valid_files)} arquivos para qualificação...")
        
        # Processar cada arquivo individualmente
        documentos_analisados = []
        textos_extraidos = []
        
        for i, file in enumerate(valid_files):
            try:
                original_filename = secure_filename(file.filename or f'arquivo_{i}.pdf')
                print(f"📄 Processando arquivo {i+1}/{len(valid_files)}: {original_filename}")
                
                # Processar arquivo de forma segura
                if Config.SECURE_PROCESSING:
                    temp_file_path, file_id = secure_manager.process_file_securely(
                        file, original_filename, user_ip
                    )
                else:
                    file_id = str(uuid.uuid4())
                    upload_filename = f"{file_id}_{original_filename}"
                    temp_file_path = os.path.join(Config.UPLOAD_FOLDER, upload_filename)
                    file.save(temp_file_path)
                
                temp_files.append(temp_file_path)
                
                # Extrair texto do PDF
                text_content = ""
                print(f"🔍 Iniciando extração de texto para {original_filename}...")
                try:
                    # Descriptografar se necessário
                    if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
                        temp_file_path = secure_manager.decrypt_file(temp_file_path)
                    
                    print(f"📄 Arquivo temporário: {temp_file_path}")
                    print(f"📁 Arquivo existe: {os.path.exists(temp_file_path)}")
                    
                    # Tentar extrair texto diretamente
                    with open(temp_file_path, 'rb') as f:
                        pdf_reader = pypdf.PdfReader(f)
                        print(f"📊 Número de páginas: {len(pdf_reader.pages)}")
                        for i, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            text_content += page_text + "\n"
                            print(f"📝 Página {i+1}: {len(page_text)} caracteres")
                    
                    print(f"✅ Extração direta: {len(text_content)} caracteres totais")
                except Exception as e:
                    print(f"❌ Erro ao extrair texto de {original_filename}: {str(e)}")
                    text_content = ""
                
                # Limpar e normalizar texto
                text_content = re.sub(r'\s+', ' ', text_content.replace('\n', ' ')).strip()
                print(f"🧹 Texto após limpeza: {len(text_content)} caracteres")
                
                # SEMPRE tentar OCR primeiro para garantir melhor extração de texto
                print(f"📄 Executando OCR para {original_filename} (SEMPRE para qualificação)...")
                try:
                    from ai.ocr_service import process_pdf_with_ocr, extract_text_from_pdf
                    temp_ocr_path = temp_file_path + '_ocr.pdf'
                    ocr_result = process_pdf_with_ocr(temp_file_path, temp_ocr_path)
                    if ocr_result.get('success'):
                        print(f"✅ OCR bem-sucedido para {original_filename}")
                        ocr_text = extract_text_from_pdf(temp_ocr_path)
                        ocr_text = re.sub(r'\s+', ' ', ocr_text.replace('\n', ' ')).strip()
                        
                        # Usar o melhor texto (OCR ou extração direta)
                        if len(ocr_text) > len(text_content):
                            text_content = ocr_text
                            print(f"✅ Usando texto do OCR ({len(ocr_text)} chars)")
                        else:
                            print(f"✅ Usando texto direto ({len(text_content)} chars)")
                    else:
                        print(f"⚠️ OCR falhou, usando extração direta")
                except Exception as ocr_error:
                    print(f"⚠️ Erro durante OCR de {original_filename}: {str(ocr_error)}")
                    print(f"✅ Continuando com extração direta")
                
                if text_content and len(text_content.strip()) > 10:
                    documentos_analisados.append({
                        'filename': original_filename,
                        'file_id': file_id,
                        'text_length': len(text_content),
                        'text_preview': text_content[:200]
                    })
                    textos_extraidos.append(f"=== DOCUMENTO: {original_filename} ===\n{text_content}\n")
                else:
                    print(f"⚠️ Texto insuficiente em {original_filename}")
                    documentos_analisados.append({
                        'filename': original_filename,
                        'file_id': file_id,
                        'text_length': 0,
                        'error': 'Texto insuficiente ou não legível'
                    })
                
            except Exception as e:
                print(f"❌ Erro ao processar {original_filename}: {str(e)}")
                documentos_analisados.append({
                    'filename': original_filename,
                    'file_id': str(uuid.uuid4()),
                    'error': str(e)
                })
        
        # Se não há textos suficientes, retornar erro
        if not textos_extraidos:
            return jsonify({
                'error': 'Não foi possível extrair texto suficiente dos documentos.',
                'details': 'Todos os arquivos podem estar vazios, corrompidos ou não conter texto legível.',
                'documentos_analisados': documentos_analisados
            }), 400
        
        # Combinar todos os textos para análise
        texto_completo = "\n\n".join(textos_extraidos)
        print(f"📝 Texto combinado: {len(texto_completo)} caracteres")
        
        # Obter modelo selecionado
        model = request.form.get('model', 'gpt-4o')
        print(f"🎯 Modelo recebido no endpoint /api/qualificacao: {model}")
        
        # Analisar com OpenAI usando nova lógica avançada
        try:
            from ai.openai_service import analyze_qualification_documents
            # Extrair nomes dos arquivos para análise
            filenames = [doc['filename'] for doc in documentos_analisados if 'filename' in doc]
            # Usar a função especializada de qualificação
            campos = analyze_qualification_documents(textos_extraidos, filenames, model=model)
        except Exception as ia_error:
            print(f"❌ Erro na análise pela IA: {ia_error}")
            return jsonify({
                'error': 'Erro ao analisar documentos com IA.',
                'details': str(ia_error),
                'documentos_analisados': documentos_analisados
            }), 500
        
        # Limpar arquivos temporários
        for temp_file in temp_files:
            if Config.SECURE_PROCESSING:
                secure_manager.cleanup_file(temp_file, user_ip)
        
        print(f"✅ REQUISIÇÃO FINALIZADA - ID: {request_id} - Sucesso!")
        return jsonify({
            'success': True,
            'message': f'Análise técnica rigorosa de qualificação concluída! {len(documentos_analisados)} documentos processados.',
            'documentos_analisados': documentos_analisados,
            'campos': campos,
            'model': model,
            'total_text_length': len(texto_completo),
            'secure_processing': Config.SECURE_PROCESSING,
            'analise_avancada': True,
            'contexto_tecnico': 'Análise realizada por IA com contexto de registrador imobiliário',
            'request_id': request_id
        })
        
    except Exception as e:
        print(f"❌ REQUISIÇÃO FALHOU - ID: {request_id} - Erro: {str(e)}")
        # Garantir limpeza em caso de erro
        for temp_file in temp_files:
            if Config.SECURE_PROCESSING:
                secure_manager.cleanup_file(temp_file, user_ip)
        return jsonify({'error': f'Erro interno do servidor: {str(e)}', 'request_id': request_id}), 500 

 

@ai_bp.route('/api/memorial', methods=['POST'])
def process_memorial():
    """Endpoint para processar arquivos DOCX de memorial de incorporação"""
    import time
    start_time = time.time()
    temp_files = []
    user_ip = request.remote_addr
    
    try:
        if 'files[]' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        
        files = request.files.getlist('files[]')
        if not files or all(file.filename == '' for file in files):
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        
        # Validar arquivos - apenas DOCX
        valid_files = []
        for file in files:
            if file.filename and file.filename.lower().endswith('.docx'):
                valid_files.append(file)
            else:
                return jsonify({'error': f'Arquivo inválido: {file.filename}. Apenas arquivos DOCX são permitidos.'}), 400
        
        if not valid_files:
            return jsonify({'error': 'Nenhum arquivo DOCX válido encontrado'}), 400
        
        print(f"📁 Processando {len(valid_files)} arquivos DOCX para memorial...")
        
        # Importar funções do extrator otimizado
        import sys
        import os
        sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        
        from extrator_memorial_otimizado import processar_arquivo_otimizado as processar_arquivo
        
        # Processar arquivos em paralelo para melhor performance
        todos_dfs = []
        documentos_processados = []
        
        # Usar ThreadPoolExecutor para processamento paralelo
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        def processar_arquivo_single(file_info):
            i, file = file_info
            try:
                original_filename = secure_filename(file.filename or f'arquivo_{i}.docx')
                
                # Processar arquivo de forma segura
                if Config.SECURE_PROCESSING:
                    temp_file_path, file_id = secure_manager.process_file_securely(
                        file, original_filename, user_ip
                    )
                else:
                    file_id = str(uuid.uuid4())
                    upload_filename = f"{file_id}_{original_filename}"
                    temp_file_path = os.path.join(Config.UPLOAD_FOLDER, upload_filename)
                    file.save(temp_file_path)
                
                # Descriptografar arquivo se necessário
                if Config.SECURE_PROCESSING and Config.ENCRYPT_TEMP_FILES:
                    temp_file_path = secure_manager.decrypt_file(temp_file_path)
                
                # Verificar se arquivo existe
                if not os.path.exists(temp_file_path):
                    return {
                        'filename': original_filename,
                        'file_id': file_id,
                        'rows_extracted': 0,
                        'error': 'Arquivo não encontrado após descriptografia',
                        'df': None,
                        'temp_file': None
                    }
                
                # Processar arquivo DOCX usando o extrator
                try:
                    df = processar_arquivo(temp_file_path)
                except Exception as process_error:
                    return {
                        'filename': original_filename,
                        'file_id': file_id,
                        'rows_extracted': 0,
                        'error': f'Erro no processamento: {str(process_error)}',
                        'df': None,
                        'temp_file': temp_file_path
                    }
                
                if df is not None and not df.empty:
                    return {
                        'filename': original_filename,
                        'file_id': file_id,
                        'rows_extracted': len(df),
                        'tipo_documento': df['Tipo Documento'].iloc[0] if 'Tipo Documento' in df.columns else 'desconhecido',
                        'df': df,
                        'temp_file': temp_file_path
                    }
                else:
                    return {
                        'filename': original_filename,
                        'file_id': file_id,
                        'rows_extracted': 0,
                        'error': 'Nenhum dado extraído',
                        'df': None,
                        'temp_file': temp_file_path
                    }
                    
            except Exception as e:
                return {
                    'filename': original_filename if 'original_filename' in locals() else f'arquivo_{i}.docx',
                    'file_id': file_id if 'file_id' in locals() else 'unknown',
                    'rows_extracted': 0,
                    'error': str(e),
                    'df': None,
                    'temp_file': temp_file_path if 'temp_file_path' in locals() else None
                }
        
        # Processar arquivos em paralelo
        with ThreadPoolExecutor(max_workers=min(4, len(valid_files))) as executor:
            future_to_file = {
                executor.submit(processar_arquivo_single, (i, file)): i 
                for i, file in enumerate(valid_files)
            }
            
            for future in as_completed(future_to_file):
                result = future.result()
                documentos_processados.append({
                    'filename': result['filename'],
                    'file_id': result['file_id'],
                    'rows_extracted': result['rows_extracted'],
                    'tipo_documento': result.get('tipo_documento', 'desconhecido'),
                    'error': result.get('error', None)
                })
                
                if result['df'] is not None:
                    todos_dfs.append(result['df'])
                
                if result['temp_file']:
                    temp_files.append(result['temp_file'])
        
        # Combinar todos os dados
        if todos_dfs:
            resultado = pd.concat(todos_dfs, ignore_index=True)
            
            # Converter para formato JSON
            dados_json = resultado.to_dict('records')
            
            # Criar arquivo Excel temporário
            excel_filename = f"memorial_{uuid.uuid4()}.xlsx"
            excel_path = os.path.join(Config.PROCESSED_FOLDER, excel_filename)
            resultado.to_excel(excel_path, index=False)
            
            # Armazenar referência do arquivo para download posterior
            memorial_files[excel_filename] = excel_path
            
            # Calcular tempo de processamento
            end_time = time.time()
            processing_time = end_time - start_time
            
            response_data = {
                'success': True,
                'message': f'Processamento concluído! {len(todos_dfs)} arquivo(s) processado(s)',
                'data': dados_json,
                'columns': resultado.columns.tolist(),
                'excel_file': excel_filename,
                'total_rows': len(resultado),
                'documentos_processados': documentos_processados,
                'processing_time': round(processing_time, 2),
                'resumo': {
                    'arquivos_processados': len(valid_files),
                    'dados_extraidos': len(resultado),
                    'tipos_encontrados': resultado['Formato'].value_counts().to_dict() if 'Formato' in resultado.columns else {}
                }
            }
            
            print(f"✅ Processamento concluído: {len(resultado)} registros extraídos")
            print(f"📊 Dados JSON criados: {len(dados_json)} registros")
            print(f"📊 Primeiro registro: {dados_json[0] if dados_json else 'N/A'}")
            return jsonify(response_data)
        else:
            return jsonify({
                'success': False,
                'error': 'Nenhum dado foi extraído dos arquivos fornecidos',
                'documentos_processados': documentos_processados
            }), 400
            
    except Exception as e:
        print(f"❌ Erro geral no processamento de memorial: {str(e)}")
        return jsonify({'error': f'Erro no processamento: {str(e)}'}), 500
    
    finally:
        # Limpeza de arquivos temporários
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                print(f"⚠️ Erro ao limpar arquivo temporário {temp_file}: {e}") 

@ai_bp.route('/api/memorial/download/<filename>', methods=['GET'])
def download_memorial_excel(filename):
    """Endpoint para download do arquivo Excel do memorial"""
    try:
        if filename not in memorial_files:
            return jsonify({'error': 'Arquivo não encontrado'}), 404
        
        file_path = memorial_files[filename]
        if not os.path.exists(file_path):
            return jsonify({'error': 'Arquivo não existe no servidor'}), 404
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        print(f"❌ Erro ao fazer download do arquivo {filename}: {str(e)}")
        return jsonify({'error': f'Erro ao fazer download: {str(e)}'}), 500

@ai_bp.route('/api/memorial/download-custom', methods=['POST'])
def download_memorial_custom():
    """Endpoint para download personalizado do Excel com colunas configuradas pelo usuário"""
    try:
        data = request.get_json()
        if not data or 'columns' not in data or 'data' not in data:
            return jsonify({'error': 'Dados inválidos'}), 400
        
        columns_config = data['columns']
        memorial_data = data['data']
        
        if not columns_config or not memorial_data:
            return jsonify({'error': 'Configuração de colunas ou dados não fornecidos'}), 400
        
        # Importar pandas e openpyxl para criar o Excel personalizado
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        import tempfile
        from datetime import datetime
        
        # Mapear IDs das colunas para campos dos dados
        column_mapping = {
            'apartamento': 'Apartamento',
            'tipo': 'Tipo', 
            'torre_bloco': 'Torre/Bloco',
            'area_privativa': 'Área Privativa (m²)',
            'area_comum': 'Área Comum (m²)',
            'area_total': 'Área Total (m²)',
            'fracao_ideal': 'Fração Ideal (%)',
            'descricao': 'Descrição'
        }
        
        # Criar DataFrame personalizado
        custom_data = []
        
        for row in memorial_data:
            custom_row = {}
            
            for col_config in columns_config:
                col_id = col_config['id']
                col_name = col_config['name']
                
                if col_config['type'] == 'custom':
                    # Coluna personalizada - usar valor padrão
                    custom_row[col_name] = col_config.get('defaultValue', '')
                else:
                    # Coluna padrão - mapear do dado original
                    original_field = column_mapping.get(col_id, col_id)
                    custom_row[col_name] = row.get(original_field, '')
            
            custom_data.append(custom_row)
        
        # Criar DataFrame
        df = pd.DataFrame(custom_data)
        
        # Criar arquivo Excel com formatação
        wb = Workbook()
        ws = wb.active
        ws.title = "Memorial Personalizado"
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # Adicionar cabeçalhos
        for col_idx, column in enumerate(df.columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=column)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Adicionar dados
        for row_idx, row in enumerate(df.values, 2):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Ajustar largura das colunas
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Salvar em arquivo temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            wb.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        # Retornar arquivo
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'memorial_personalizado_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        print(f"❌ Erro ao gerar Excel personalizado: {str(e)}")
        return jsonify({'error': f'Erro ao gerar arquivo personalizado: {str(e)}'}), 500 